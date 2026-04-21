from pathlib import Path
from .base_loader import BaseLoader


class DocxLoader(BaseLoader):
    """Load DOCX documents dùng python-docx (không phụ thuộc LlamaIndex API)."""

    def load(self, file_path: Path) -> list[str]:
        """Đọc toàn bộ text từ DOCX, ghép các paragraph thành chuỗi."""
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return ["\n".join(paragraphs)] if paragraphs else []

    def load_pages(self, file_path: Path) -> list[tuple[int, str]]:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Chia thành "pages" giả theo 20 paragraph mỗi block
        pages = []
        block_size = 20
        for i in range(0, len(paragraphs), block_size):
            block = "\n".join(paragraphs[i:i + block_size])
            pages.append((i // block_size + 1, block))
        return pages
